from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
OUT = Path(__file__).resolve().parent
LOGO = ROOT / "assets" / "Stars - Original Logo.png"
SHOTS = OUT / "screenshots"
ANNOTATED = OUT / "screenshots_annotated"
FONT = "Tahoma"
NAVY = "0B2545"
BLUE = "0B84C6"
LIGHT_BLUE = "E8F3FA"
LIGHT_GRAY = "F2F4F7"
GREEN = "E5F7EF"
AMBER = "FFF4D6"
RED = "FDECEC"
MUTED = "64748B"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def setup_doc(title, subtitle, audience, version="0.2"):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for name, size, color, before, after in [
        ("Heading 1", 17, NAVY, 16, 8),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.19)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.2

    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"RATS COMMAND CENTER | {audience}")
    set_font(r, 8.5, True, MUTED)
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RATS Command Center v0.2  •  หน้า ")
    set_font(r, 8.5, False, MUTED)
    add_page_field(p)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture = p.add_run().add_picture(str(LOGO), width=Inches(1.45))
        doc_pr = picture._inline.docPr
        doc_pr.set("descr", "โลโก้บริษัท STARS Microelectronics")
        doc_pr.set("title", "STARS Microelectronics")
        p.paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 27, True, NAVY)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_font(r, 15, False, BLUE)
    p.paragraph_format.space_after = Pt(30)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"สำหรับ {audience}\nเวอร์ชันระบบ {version} | ปรับปรุง 28 สิงหาคม 2569")
    set_font(r, 11, False, MUTED)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(55)
    callout(doc, "วัตถุประสงค์", "เอกสารนี้อธิบายการใช้งานตามระบบที่ติดตั้งจริง ควรใช้ร่วมกับขั้นตอนความปลอดภัยของโรงงานและสิทธิ์ที่ได้รับมอบหมาย", "info")
    doc.add_page_break()
    return doc


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_font(r)


def callout(doc, label, text, kind="info"):
    fill = {"info": LIGHT_BLUE, "warn": AMBER, "danger": RED, "ok": GREEN}[kind]
    table = doc.add_table(rows=1, cols=1)
    mark_header_row(table.rows[0])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.55)
    cell = table.cell(0, 0)
    cell.width = Inches(6.55)
    shade(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, 10.5, True, NAVY)
    r = p.add_run(text)
    set_font(r, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def figure(doc, filename, caption, alt, width=6.35):
    path = ANNOTATED / filename
    if not path.exists():
        path = SHOTS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", alt)
    picture._inline.docPr.set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = True
    r = cp.add_run(caption)
    set_font(r, 9, True, MUTED)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    mark_header_row(tbl.rows[0])
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    if widths is None:
        widths = [6.5 / len(headers)] * len(headers)
    for i, (cell, head) in enumerate(zip(tbl.rows[0].cells, headers)):
        cell.width = Inches(widths[i])
        shade(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(head)
        set_font(r, 9.5, True, "FFFFFF")
    for row_i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_i % 2:
                shade(cells[i], LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, 9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tbl


def toc(doc, items):
    heading(doc, "สารบัญ", 1)
    para(doc, "สารบัญนี้เรียงตามหัวข้อหลักของคู่มือ เพื่อให้ค้นหาขั้นตอนได้รวดเร็ว")
    for no, title in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(title)
        set_font(r, 11)
    doc.add_page_break()


def user_manual():
    doc = setup_doc(
        "คู่มือการใช้งาน RATS Command Center",
        "RATS Recipe Automation & Transfer System",
        "ผู้ใช้งานทั่วไป พนักงานปฏิบัติการ และช่างเทคนิค",
    )
    toc(doc, [
        (1, "รู้จักระบบและขอบเขตการใช้งาน"), (2, "การเข้าสู่ระบบและสิทธิ์"),
        (3, "ทำความเข้าใจหน้าจอ Dashboard"), (4, "การเลือกและตรวจสอบเครื่องจักร"),
        (5, "การดึงสูตรจากเครื่อง (Pull)"), (6, "การส่งสูตรไปเครื่อง (Push)"),
        (7, "การลบสูตรจากเครื่อง"), (8, "การอัปเดตสูตร PPID ซ้ำแบบอัตโนมัติ"),
        (9, "เหตุการณ์จาก Recipe Bot"), (10, "บันทึกเหตุการณ์และการแจ้งเตือน"),
        (11, "การแก้ปัญหาเบื้องต้น"), (12, "เช็กลิสต์ปฏิบัติงาน"),
    ])

    heading(doc, "1. รู้จักระบบและขอบเขตการใช้งาน")
    para(doc, "RATS Command Center เป็นหน้าจอกลางสำหรับดูสถานะเครื่อง Wire Bonder และจัดการ Recipe ระหว่าง Host กับเครื่องจักร โดยการสั่งงานเครื่องใช้ SECS/GEM ส่วน Recipe Bot ใช้รับไฟล์ใหม่หรือไฟล์ที่ช่างแก้ไขจากเครื่องกลับมายัง Host")
    table(doc, ["องค์ประกอบ", "หน้าที่"], [
        ("Dashboard", "เลือกเครื่อง ดูสถานะ สั่ง Pull/Push/Delete และดู Event Log"),
        ("Python Backend", "รับคำสั่งจาก Dashboard ตรวจสิทธิ์ และกระจายสถานะแบบเรียลไทม์"),
        ("Machine Link", "ช่องทาง SECS/GEM ระหว่าง Server กับเครื่องจักร"),
        ("Recipe Bot", "ตรวจจับไฟล์ PWB ใหม่/เปลี่ยนบนเครื่อง และส่งกลับ Host หลังผู้ใช้ยืนยัน"),
    ], [1.55, 4.95])
    callout(doc, "ข้อควรจำ", "สถานะเครื่องจักรและสถานะ Recipe Bot เป็นคนละสถานะ เครื่องอาจออนไลน์แต่ Bot ออฟไลน์ หรือกลับกันได้", "warn")

    heading(doc, "2. Procedure: การเข้าสู่ระบบและสิทธิ์")
    para(doc, "เปิดเบราว์เซอร์ไปยัง URL ที่ฝ่าย IT หรือ Developer แจ้ง เช่น http://<IP-Server>:3000 จากนั้นกดเข้าสู่ระบบ กรอก Employee Number ของตนเอง พร้อมชื่อผู้ใช้และรหัสผ่านที่ได้รับอนุญาต ไม่ควรใช้ Employee Number หรือบัญชีของผู้อื่น")
    figure(doc, "01_login_annotated.png", "รูปที่ 1 หน้าต่าง Login แบบระบุตำแหน่ง", "ภาพครอปหน้าต่างเข้าสู่ระบบพร้อมลูกศรหมายเลข 1 ถึง 4 ชี้ Employee Number ชื่อผู้ใช้ รหัสผ่าน และปุ่มเข้าสู่ระบบ", width=5.0)
    table(doc, ["หมายเลข", "ตำแหน่ง", "สิ่งที่ต้องทำ"], [
        ("1", "Employee Number", "กรอกรหัสพนักงานของผู้ปฏิบัติงาน ห้ามใช้รหัสของผู้อื่น"),
        ("2", "ชื่อผู้ใช้", "กรอกบัญชีตาม Role ที่ได้รับมอบหมาย"),
        ("3", "รหัสผ่าน", "กรอกรหัสผ่านของ Role โดยไม่เปิดเผยให้ผู้อื่น"),
        ("4", "เข้าสู่ระบบ", "กดหลังตรวจทั้งสามช่องครบ แล้วตรวจ EN/Role มุมขวาบน"),
    ], [0.75, 1.5, 4.25])
    bullets(doc, [
        "ตรวจ URL ให้เป็น Server ของหน่วยงาน ไม่ใช่เว็บภายนอก",
        "กด “เข้าสู่ระบบ / ยืนยันสิทธิ์”",
        "กรอก Employee Number ของตนเอง ช่องนี้ห้ามเว้นว่าง",
        "กรอกชื่อผู้ใช้และรหัสผ่านที่ได้รับมอบหมาย แล้วกด “เข้าสู่ระบบ”",
        "ตรวจมุมขวาบนว่าแสดง EN และ Role ของตนถูกต้อง",
    ], numbered=True)
    table(doc, ["สิทธิ์", "สิ่งที่ทำได้โดยทั่วไป"], [
        ("Guest", "ยังไม่สามารถเข้าหน้า RATS ได้ ต้องเข้าสู่ระบบ"),
        ("Operator", "ดู Dashboard เลือกเครื่อง ตรวจสถานะ และใช้งานมาตรฐาน Pull/Push"),
        ("Technician", "สิทธิ์ Operator รวมทั้ง Delete และล้าง Event Log"),
        ("Administrator", "สิทธิ์ Technician และดู Employee Activity Audit"),
        ("Developer", "สิทธิ์สูงสุด รวมถึง Deploy ไฟล์ Recipe Bot/Config ไปยังเครื่อง"),
    ], [1.45, 5.05])
    para(doc, "ระบบจะเตือนก่อนออกจากระบบเมื่อไม่มีการใช้งานครบช่วงเวลาที่กำหนด (ค่าเริ่มต้น 5 นาที) หากยังทำงานอยู่ให้กด “คงสถานะการเข้าสู่ระบบ” มิฉะนั้นระบบจะกลับเป็น Guest โดยอัตโนมัติ")

    heading(doc, "3. Procedure: ทำความเข้าใจหน้าจอ Dashboard")
    figure(doc, "02_operator_wb83_annotated.png", "รูปที่ 2 แผงควบคุมเครื่องที่เลือกพร้อมตำแหน่งสำคัญ", "ภาพครอป Dashboard Operator พร้อมลูกศรหมายเลข 1 ถึง 5 ชี้ชื่อเครื่อง ปุ่ม Pull ปุ่ม Push ช่อง Recipe และ Event Log")
    table(doc, ["หมายเลข", "ตำแหน่ง", "วิธีอ่าน/ใช้งาน"], [
        ("1", "เครื่องที่เลือก", "ตรวจ WB# ให้ตรงกับป้ายบนเครื่องจริงก่อนทุกคำสั่ง"),
        ("2", "Pull", "ดึง/ซิงก์สูตรจากเครื่องไปยัง Host ผ่าน SECS/GEM"),
        ("3", "Push", "ส่งสูตรจาก Host ไปยังเครื่อง หลังเลือก Recipe และตรวจสถานะ"),
        ("4", "Recipe selector", "เลือกหรือพิมพ์ PPID และตรวจตัวอักษร ขีดล่าง ตัวเลขให้ครบ"),
        ("5", "Event Log", "อ่าน SUCCESS/ALERT/ERROR และเวลาเกิดเหตุ ห้ามดูเฉพาะสีสถานะ"),
    ], [0.75, 1.5, 4.25])
    table(doc, ["สถานะ", "ความหมาย", "การปฏิบัติ"], [
        ("ONLINE", "ช่องทางนั้นเชื่อมต่อและตอบสนอง", "ทำงานต่อได้ แต่ตรวจเครื่องจริงก่อนคำสั่งสำคัญ"),
        ("OFFLINE", "ขาดการเชื่อมต่อหรือไม่ตอบสนอง", "หยุด Push/Delete และแจ้งช่าง/Developer"),
        ("CONNECTING", "ระบบกำลังเชื่อมต่อ", "รอสั้น ๆ และดู Event Log"),
        ("UNCHECKED", "ยังไม่มีผลตรวจยืนยัน", "เลือกเครื่องหรือทดสอบการเชื่อมต่อ"),
    ], [1.25, 2.35, 2.9])

    heading(doc, "4. Procedure: การเลือกและตรวจสอบเครื่องจักร")
    bullets(doc, [
        "เลือกจากรายการ Wire Bonder ทางซ้าย หรือสแกน Barcode/Serial ในช่องค้นหา",
        "ตรวจชื่อเครื่องและรหัส WB# ให้ตรงกับเครื่องจริงทุกครั้ง",
        "ตรวจสถานะ Machine และ Recipe Bot แยกกัน",
        "กด “ทดสอบการเชื่อมต่อ” เมื่อไม่แน่ใจ แล้วอ่านผลด้านบนหรือ Event Log",
    ], numbered=True)
    figure(doc, "02_operator_wb83_annotated.png", "รูปที่ 3 ตัวอย่างเลือก WB#83 และตรวจตำแหน่งก่อนสั่งงาน", "Dashboard หลังเลือก Wire Bonder 83 พร้อมหมายเลขชี้ตำแหน่งตรวจสอบ")
    callout(doc, "ก่อนสั่งงาน", "อย่าอาศัยคำว่า ONLINE อย่างเดียว ต้องยืนยันเครื่องที่เลือกและชื่อ Recipe ให้ถูกต้อง เพราะ Push/Delete มีผลกับเครื่องจริง", "danger")

    heading(doc, "5. Procedure: การดึงสูตรจากเครื่อง (Pull)")
    para(doc, "Pull ใช้ซิงก์รายการ/สูตรจากเครื่องผ่าน SECS/GEM ไปยัง Host ตามความสามารถของเครื่องและ Section Manager")
    bullets(doc, [
        "เลือกเครื่องและยืนยันว่า Machine Status เป็น ONLINE",
        "เข้าสู่สิทธิ์ Operator ขึ้นไป",
        "กด “ดึงสูตรการผลิต (Pull)” และรอจนสถานะ SYNCING สิ้นสุด",
        "ตรวจผล SUCCESS/ERROR และ Event Log รวมถึงรายการที่ดึงได้หรือถูกข้าม",
    ], numbered=True)
    callout(doc, "ห้ามทำ", "อย่าปิด Server หรือรีสตาร์ตเครื่องระหว่างกำลัง Pull หากไม่จำเป็น", "warn")

    heading(doc, "6. Procedure: การส่งสูตรไปเครื่อง (Push)")
    figure(doc, "02_operator_wb83_annotated.png", "รูปที่ 4 ตำแหน่ง Push และ Recipe selector", "หมายเลข 3 ชี้ปุ่ม Push และหมายเลข 4 ชี้ช่องเลือก Recipe ก่อนส่งไปเครื่อง")
    bullets(doc, [
        "เลือกเครื่องปลายทางและตรวจ Machine Status = ONLINE",
        "เลือก Recipe จากรายการ หรือพิมพ์ชื่อ PPID ที่ต้องการ",
        "ตรวจชื่อซ้ำอีกครั้ง โดยเฉพาะขีดล่าง ตัวเลข และคำต่อท้าย",
        "กด “ส่งสูตรการผลิต (Push)” ระบบอาจเสนอชื่อใกล้เคียงหากไม่พบชื่อที่พิมพ์",
        "หากชื่อที่เสนอถูกต้องให้ยืนยัน หากไม่ถูกต้องให้ยกเลิก",
        "รอผลสำเร็จและตรวจ Recipe ปัจจุบันบน Dashboard/เครื่องจริง",
    ], numbered=True)
    para(doc, "การ Push ไปเครื่องใช้ SECS/GEM เพื่อให้ Recipe ปรากฏและถูกจัดการใน HSMS/GEM ของเครื่อง ไม่ใช่การคัดลอกไฟล์ธรรมดา")

    heading(doc, "7. Procedure: การลบสูตรจากเครื่อง")
    bullets(doc, [
        "เลือกเครื่องและ Recipe ที่ต้องการลบอย่างระมัดระวัง",
        "กดไอคอนถังขยะ ระบบจะแสดงหน้าต่างยืนยัน",
        "อ่านชื่อ Recipe และชื่อเครื่องในหน้าต่างยืนยัน",
        "กด “ยืนยันการลบ” เฉพาะเมื่อแน่ใจ แล้วตรวจผลใน Event Log",
    ], numbered=True)
    callout(doc, "คำเตือนสำคัญ", "Delete ลบ Recipe จากเครื่องและไม่สามารถย้อนกลับจากหน้าจอนี้ได้ ควรมีสำเนาบน Host และได้รับอนุญาตก่อนเสมอ", "danger")

    heading(doc, "8. Procedure: การอัปเดตสูตร PPID ซ้ำแบบอัตโนมัติ")
    para(doc, "เมื่อ Recipe Bot ส่งไฟล์ที่มี PPID เดียวกับไฟล์บน Host แต่เนื้อหาแตกต่าง Backend จะเก็บไฟล์ Host เดิมไว้ใน .archive แล้วเขียนทับด้วยไฟล์ใหม่โดยอัตโนมัติ ไม่ต้องรอผู้ใช้อนุมัติบน Dashboard")
    bullets(doc, [
        "ผู้ใช้ที่เครื่องยืนยันส่งสูตรจาก Popup Recipe Ready",
        "Bot ส่ง Snapshot ไป Host ผ่านช่องทาง TCP 5003",
        "Backend เปรียบเทียบ PPID และ Hash; ถ้าเนื้อหาเปลี่ยนจะ Archive ไฟล์เดิมและติดตั้งไฟล์ใหม่",
        "ตรวจข้อความ SUCCESS ใน Event Log และ Popup Transfer Complete ที่เครื่อง",
    ], numbered=True)
    callout(doc, "หลักการ", "PPID เหมือนกันไม่ได้แปลว่าไฟล์เหมือนกัน ระบบเปรียบเทียบ Hash และเก็บสำเนาเดิมใน .archive ก่อนเขียนทับเสมอ", "info")

    heading(doc, "9. Procedure: เหตุการณ์จาก Recipe Bot บนเครื่อง")
    para(doc, "Recipe Bot เฝ้าดูไฟล์ NPGM0.PWB ถึง NPGMn.PWB ในโฟลเดอร์ BONDPROG แบบ Event-driven เมื่อพบไฟล์ใหม่หรือไฟล์เปลี่ยน จะอ่าน Program Name ภายในไฟล์เป็น PPID จริง")
    figure(doc, "17_recipe_bot_recipe_ready.jpg", "รูปที่ 5 Popup Recipe Ready หลังเครื่องบันทึก Bond Program", "Recipe Bot แสดง PPID และชื่อไฟล์ NPGM บนเครื่อง พร้อมปุ่ม Yes/No เพื่อยืนยันการส่งไป Factory Host", width=4.15)
    bullets(doc, [
        "ไฟล์เหมือนกับ Host ทุกประการ: Bot ข้ามโดยเงียบเพื่อไม่ให้เกิด Popup วนซ้ำ",
        "ไฟล์ใหม่/เปลี่ยน: Operator ที่เครื่องจะเห็น Popup ขออนุญาตส่ง",
        "กดยอมรับ: Bot เก็บ Snapshot และส่งไป Host เมื่อช่องทางพร้อม หากเน็ตขาดจะค้างใน recipe_outbox เพื่อ Retry",
        "เมื่อ Host ยืนยันรับสำเร็จ: Popup สำเร็จจะค้างจนผู้ใช้กด OK; การกด OK ปิดเฉพาะ Popup ไม่ปิด Bot",
    ])
    figure(doc, "16_recipe_bot_transfer_complete.jpg", "รูปที่ 6 Popup Transfer Complete หลัง Host รับสูตรสำเร็จ", "Recipe Bot แจ้ง PPID ที่ส่งสำเร็จ ผู้ใช้ต้องกด OK เพื่อรับทราบ โดย Bot ยังคงทำงานต่อ", width=4.15)

    heading(doc, "10. Procedure: บันทึกเหตุการณ์และการแจ้งเตือน")
    para(doc, "Event Log แสดงเฉพาะเหตุการณ์ล่าสุดและอัปเดตผ่าน WebSocket เมื่อสถานะเปลี่ยน ไม่ใช้การ Poll เครื่องอย่างต่อเนื่อง")
    table(doc, ["ระดับ", "ใช้กับเหตุการณ์"], [
        ("SUCCESS", "เชื่อมต่อสำเร็จ รับ/ส่ง/ติดตั้งไฟล์สำเร็จ"),
        ("INFO", "เริ่มคำสั่ง เปลี่ยนสถานะ หรือข้อมูลทั่วไป"),
        ("ALERT/WARN", "ช่องทาง Offline, ไฟล์รอเปลี่ยน, หรือเหตุการณ์ที่ต้องตรวจ"),
        ("ERROR", "คำสั่งล้มเหลว ติดต่อไม่ได้ หรือข้อมูลไม่ถูกต้อง"),
    ], [1.5, 5.0])
    para(doc, "การล้าง Log ต้องใช้สิทธิ์ Technician ขึ้นไป Backend จะ Archive Log ลงไฟล์ก่อนล้างออกจากหน่วยความจำ")
    para(doc, "ระบบบันทึก Employee Number, Role, เวลา Login/Logout, เหตุผลการออก และ Action เช่น Pull/Push พร้อม Machine/PPID โดย Administrator หรือ Developer ตรวจได้ในหน้า System Status")
    figure(doc, "15_developer_employee_audit.png", "รูปที่ 7 หน้า Employee Activity Audit", "หน้า System Status แสดง Employee Number, Role, เวลาเข้า-ออก และรายการ Action ของแต่ละ Session")

    heading(doc, "11. การแก้ปัญหาเบื้องต้น")
    table(doc, ["อาการ", "ตรวจตามลำดับ"], [
        ("Dashboard เปิดไม่ได้", "ตรวจ URL/เครือข่าย แล้วแจ้งว่าเข้า Port 3000 ไม่ได้"),
        ("Backend Offline", "แจ้ง Developer ให้ตรวจ Backend Port 8080; ห้ามกดคำสั่งซ้ำรัว ๆ"),
        ("Machine Online แต่ Bot Offline", "ตรวจว่า Recipe Bot ทำงานบนเครื่อง, config.ini ถูกต้อง และ Port 5003 เปิด"),
        ("Bot Online แต่ Machine Offline", "ตรวจ SECS/GEM/HSMS ของเครื่องและ Port 5001"),
        ("Push/Delete ล้มเหลว", "ตรวจสิทธิ์ ชื่อ Recipe สถานะ Machine และข้อความ Error ใน Log"),
        ("Session หมดอายุ", "เข้าสู่ระบบใหม่ แล้วทำรายการอีกครั้งหลังตรวจสถานะล่าสุด"),
        ("ไฟล์ใหม่ไม่ถึง Host", "ตรวจ Popup ที่เครื่อง, recipe_outbox, Recipe Bot Status และแจ้งชื่อเครื่อง/เวลา/PPID"),
    ], [2.05, 4.45])
    callout(doc, "ข้อมูลที่ต้องแจ้ง", "ชื่อเครื่อง WB#, เวลาเกิดเหตุ, คำสั่งที่กด, PPID, สถานะ Machine/Bot และข้อความ Error แบบเต็ม ช่วยให้แก้ปัญหาได้เร็วที่สุด", "ok")

    heading(doc, "12. เช็กลิสต์ปฏิบัติงาน")
    heading(doc, "ก่อนเริ่มงาน", 2)
    bullets(doc, ["เข้าสู่ระบบด้วยบัญชีตนเอง", "เลือกเครื่องถูกต้อง", "Machine/Bot อยู่ในสถานะที่งานนั้นต้องใช้", "มีสำเนา Recipe และได้รับอนุญาต"])
    heading(doc, "หลังทำรายการ", 2)
    bullets(doc, ["เห็นผล SUCCESS", "ตรวจ Recipe/เครื่องจริง", "ตรวจ Event Log", "ออกจากระบบเมื่อเสร็จงาน"])
    heading(doc, "ช่องทางส่งต่อปัญหา", 2)
    para(doc, "หากแก้ตามคู่มือนี้แล้วยังไม่สำเร็จ ให้หยุดการลองซ้ำที่อาจกระทบเครื่อง และส่งข้อมูลเหตุการณ์ให้ Developer/IT ตามขั้นตอนของหน่วยงาน")
    heading(doc, "Procedure: ออกจากระบบ", 2)
    bullets(doc, ["กดไอคอนออกจากระบบด้านขวาบน", "ตรวจ Role ที่กำลังใช้งาน", "กด “ยืนยันออกจากระบบ”", "ตรวจว่าหน้าจอกลับเป็น Guest"], numbered=True)
    callout(doc, "ข้อควรระวัง", "ตรวจ Employee Number และ Role มุมขวาบนก่อนยืนยันทุกครั้ง เพื่อไม่ให้ออกจากบัญชีของผู้ปฏิบัติงานคนอื่น")
    return doc


def dev_manual():
    doc = setup_doc(
        "คู่มือ Developer และผู้ดูแลระบบ",
        "RATS Command Center / RATS Recipe Transfer Platform",
        "Developer, System Administrator และผู้ดูแลการเชื่อมต่อเครื่องจักร",
    )
    toc(doc, [
        (1, "ภาพรวมสถาปัตยกรรม"), (2, "โครงสร้าง Source Code"), (3, "Network และ Port"),
        (4, "ติดตั้ง Runtime และ Build"), (5, "ตั้งค่า Server"), (6, "ฐานข้อมูลเครื่องและการเพิ่มเครื่อง"),
        (7, "Section Manager และ SECS/GEM"), (8, "Recipe Bot Port 5003"),
        (9, "Deployment Receiver Port 5004"), (10, "Authentication และ Role"),
        (11, "Recipe Workflow และการเก็บไฟล์"), (12, "Realtime State และ WebSocket"),
        (13, "Start/Stop/Log"), (14, "Test และ Acceptance"), (15, "Troubleshooting"),
        (16, "Security Hardening และ Change Control"), (17, "Production Readiness และ Release Gate"),
    ])

    heading(doc, "1. ภาพรวมสถาปัตยกรรม")
    para(doc, "ระบบใช้สถาปัตยกรรม Host เป็นฝ่ายเริ่มการเชื่อมต่อไปยังเครื่องจักร เพื่อลดปัญหาเครื่องส่ง HTTP กลับ Host และแยกช่องทางตามหน้าที่")
    callout(doc, "Data Flow", "Browser → Frontend :3000 → Backend :8080/WebSocket → Machine SECS/GEM :5001, Recipe Bot :5003 และ Deployment Receiver :5004", "info")
    table(doc, ["ชั้นระบบ", "เทคโนโลยี", "ความรับผิดชอบ"], [
        ("Frontend", "React + Vite + Tailwind", "UI, Role gating, ภาษา TH/EN, WebSocket state"),
        ("RATS Backend", "FastAPI/Uvicorn", "Auth, API, Event Log, Recipe storage, TCP client"),
        ("Section Manager", "Python + secsgem", "Worker ต่อเครื่อง, SECS/GEM Pull/Push/Delete"),
        ("Recipe Bot", "C++ Win32 32-bit static", "Watch PWB, PPID extraction, Popup, Outbox, TCP 5003"),
        ("Deployment Receiver", "C++ Win32 32-bit static", "รับเฉพาะ EXE/config ผ่าน TCP 5004 แบบ Auth/CRC32"),
    ], [1.45, 1.85, 3.2])

    heading(doc, "2. โครงสร้าง Source Code")
    table(doc, ["ตำแหน่ง", "เนื้อหา"], [
        ("arc-system/client-shell", "Frontend Dashboard; source ใน src และผล Build ใน dist"),
        ("arc-system/client-rats/main.py", "FastAPI Backend และ Host-side TCP channels"),
        ("arc-system/section-manager", "Manager/Worker สำหรับแต่ละเครื่อง"),
        ("database.py", "Single source of truth สำหรับ Machine ID, IP, Port และ Serial mapping"),
        ("BondingProg", "คลัง Recipe .PWB ของ Host รวม .pending และ .archive"),
        ("secs_proxy_bot", "Source/Build/Package ของ Recipe Bot"),
        ("tcp_deploy_bot", "Source/Build/Package ของ Deployment Receiver"),
        ("production-checklist.md", "Checklist การติดตั้ง Production"),
    ], [2.45, 4.05])

    heading(doc, "3. Network และ Port")
    table(doc, ["Port", "ทิศทาง", "บริการ", "Firewall"], [
        ("3000/TCP", "Operator PC → Server", "Frontend", "Inbound Server"),
        ("8080/TCP", "Browser → Server", "API + WebSocket", "Inbound Server"),
        ("5001/TCP", "Server → Machine", "Native SECS/GEM/HSMS", "Inbound Machine จาก Server"),
        ("5003/TCP", "Server → Machine", "Recipe Bot full-duplex", "Inbound Machine จาก Server"),
        ("5004/TCP", "Server → Machine", "Deployment Receiver", "Inbound Machine จาก Server"),
    ], [0.8, 1.8, 2.15, 1.75])
    bullets(doc, [
        "Port 5002 ไม่ได้ใช้งานและไม่จำเป็นในสถาปัตยกรรมปัจจุบัน",
        "Recipe Bot ไม่ POST HTTP ไป Port 8080; PWB กลับผ่าน Session 5003 เดิม",
        "จำกัด Source IP ของ 5003/5004 ให้เป็น RATS Server เมื่อ Firewall รองรับ",
        "หาก IP เปลี่ยนจาก 169.254.13.xx เป็น 192.168.11.xx ให้แก้ database.py และตรวจ Routing/Firewall ใหม่",
    ])

    heading(doc, "4. ติดตั้ง Runtime และ Build")
    para(doc, "ข้อกำหนดแนะนำ: Python 3.11+, Node.js 18+, npm, และ Dependencies ใน requirement.txt")
    bullets(doc, [
        "Root: pip install -r requirement.txt",
        "Frontend: cd arc-system\\client-shell แล้ว npm install",
        "Production UI: npm run build",
        "ตรวจว่า dist/index.html และ dist/assets ถูกสร้างใหม่",
        "C++ Bot: ใช้ MSYS2 MinGW32 และ mingw32-make ตาม README ของแต่ละ Bot",
    ], numbered=True)
    callout(doc, "Production", "อย่าใช้ Vite dev server เป็นบริการถาวร ควรใช้ Build ที่ตรวจแล้วและวิธี Serve ที่องค์กรอนุมัติ", "warn")

    heading(doc, "5. ตั้งค่า Server")
    table(doc, ["Environment Variable", "ค่า/หน้าที่"], [
        ("RATS_HOST", "0.0.0.0 เมื่อต้องรับจาก LAN"),
        ("RATS_PORT", "8080"),
        ("RATS_RELOAD", "false ใน Production"),
        ("RATS_CORS_ORIGINS", "รายการ Frontend Origin เช่น http://<SERVER_IP>:3000"),
        ("RATS_SESSION_TIMEOUT_MINUTES", "อายุ Session; ค่าเริ่มต้น 5 นาที"),
        ("RATS_PROXY_UPLOAD_TOKEN", "ต้องตรงกับ file_channel_token ของ Recipe Bot"),
        ("RATS_DEPLOY_TOKEN", "ต้องตรงกับ deploy_token ของ Deployment Receiver"),
        ("RATS_DEPLOY_PORT", "5004"),
        ("RATS_*_PASSWORD", "Password ของ Operator/Technician/Admin/Developer"),
    ], [2.75, 3.75])
    callout(doc, "ความปลอดภัย", "ค่า Password/Token เริ่มต้นมีไว้สำหรับ Pilot เท่านั้น ต้องเปลี่ยนก่อนใช้งานจริงและห้ามบันทึก Secret ลง Git", "danger")

    heading(doc, "6. ฐานข้อมูลเครื่องและการเพิ่มเครื่อง")
    para(doc, "แก้ MACHINE_DB ใน database.py โดยใช้ WB#nn เป็น Key และกำหนด name, ip, port=5001, bot_file_port=5003, deploy_port=5004 และ session_id=0 ให้ครบทุกเครื่องที่รองรับ Recipe Bot")
    bullets(doc, [
        "เพิ่ม Serial/Barcode mapping ใน SERIAL_TO_MACHINE",
        "ตรวจว่าเลข WB# สอดคล้องกับ IP/Asset จริง",
        "Restart Backend และ Section Manager หลังแก้ database.py",
        "ทดสอบ Lookup, Machine Link, Bot Link และ Event Log",
        "ห้ามมี Machine ID ซ้ำหรือ IP ซ้ำโดยไม่ตั้งใจ",
    ], numbered=True)

    heading(doc, "7. Section Manager และ SECS/GEM")
    para(doc, "Section Manager สร้าง Worker แยกต่อเครื่องจาก MACHINE_DB และรายงาน state กลับ Backend ผ่าน internal connection status การเชื่อมต่อหลักใช้ Native HSMS/SECS/GEM Port 5001")
    table(doc, ["คำสั่ง", "SECS/GEM ที่เกี่ยวข้อง", "ผลที่คาดหวัง"], [
        ("Test Link", "S1F1/S1F2 หรือกลไก Link ของ Worker", "ONLINE/OFFLINE"),
        ("Pull", "Recipe list/request flow", "Host ซิงก์ Recipe ใหม่และรายงาน skipped"),
        ("Push", "Program download flow", "Machine ตอบ ACKC7 และโหลด Recipe"),
        ("Delete", "S7F17/S7F18", "ลบ PPID ที่เครื่องยืนยัน"),
    ], [1.2, 2.65, 2.65])
    callout(doc, "ข้อควรระวัง", "Machine Link ONLINE ไม่ได้ยืนยันว่า Recipe Bot ONLINE เพราะเป็นคนละ Process และคนละ Port", "warn")

    heading(doc, "8. Recipe Bot Port 5003")
    para(doc, "Recipe Bot เป็น Listener บนเครื่อง แต่ RATS Host เป็นฝ่าย Connect และ Authenticate ด้วย ARCFBOT1 + Token จากนั้นใช้ช่องทางเดียวแบบ Full-duplex")
    table(doc, ["config.ini key", "คำแนะนำ"], [
        ("file_listen_ip", "0.0.0.0"), ("file_listen_port", "5003"),
        ("file_channel_token", "ต้องตรง RATS_PROXY_UPLOAD_TOKEN"),
        ("machine_id", "ควรตรง WB# ของเครื่อง; Backend ใช้ IP/Channel ที่ Host เลือกเป็นตัวตนหลักเพื่อรองรับ Config รุ่นเก่า"),
        ("max_file_bytes", "ค่าเริ่มต้น 20 MiB"),
        ("watch_dir", "C:\\SYSTEM\\BONDPROG หรือ Path จริงของเครื่อง"),
        ("file_ext", ".PWB"), ("log_file", "recipe_bot.log"),
    ], [2.35, 4.15])
    bullets(doc, [
        "ReadDirectoryChangesW ตรวจ Added/Changed แบบ Event-driven",
        "รอไฟล์ Stable ก่อน Snapshot และอ่าน Program Name เป็น PPID",
        "ตรวจ Host ก่อน Popup: new / identical / different",
        "Outbox เป็น FIFO และลบงานเมื่อ Host Confirm เท่านั้น",
        "TCP keepalive ตรวจ Disconnect; Event Log กระจายเฉพาะ State transition",
    ])
    figure(doc, "17_recipe_bot_recipe_ready.jpg", "รูปที่ D1 Popup Recipe Ready บนเครื่องจริง", "หลังบันทึก Bond Program Bot อ่าน PPID ภายใน NPGM(n).PWB และขอให้ผู้ใช้ยืนยันส่งไป Host", width=4.15)
    figure(doc, "16_recipe_bot_transfer_complete.jpg", "รูปที่ D2 Popup Transfer Complete บนเครื่องจริง", "Host ยืนยันรับไฟล์แล้ว Popup ค้างจนผู้ใช้กด OK โดยไม่ปิด Process ของ Recipe Bot", width=4.15)

    heading(doc, "9. Deployment Receiver Port 5004")
    para(doc, "ติดตั้ง Receiver ครั้งแรกด้วยวิธี Bootstrap ที่องค์กรรับรอง เช่น RDP/USB/Software deployment จากนั้น Dashboard สามารถส่ง secs_proxy_bot.exe และ config.ini ได้เฉพาะ Role Developer")
    table(doc, ["การป้องกัน", "รายละเอียด"], [
        ("Authentication", "ARCDEP01 + deploy_token"),
        ("Machine validation", "machine_id=AUTO ตรวจ WB#nn กับเลขท้าย IP ของ Interface ที่รับ Connection"),
        ("Allowlist", "รับเฉพาะ secs_proxy_bot.exe และ config.ini"),
        ("Integrity", "ตรวจ Size, CRC32 และเขียนแบบ Atomic"),
        ("Execution", "Receiver ไม่ Run ไฟล์ที่รับ"),
        ("Locked EXE", "เก็บเป็น secs_proxy_bot.exe.pending ให้หยุด Bot แล้ว Replace"),
    ], [2.0, 4.5])
    callout(doc, "ตำแหน่งไฟล์", "แนะนำให้ Receiver อยู่ C:\\ARCDeployReceiver ส่วนไฟล์ Recipe Bot ถูกส่งไป C:\\ARCRecipeBot ตาม deploy_dir", "info")
    heading(doc, "Procedure: Deploy Recipe Bot/Config จาก Dashboard", 2)
    figure(doc, "05_developer_wb83_clean.png", "รูปที่ D3 Dashboard Role Developer ขณะ WB#83 ออนไลน์และแสดง TCPBOT", "Dashboard Developer จากเครื่องจริง แสดง Machine และ Recipe Bot ของ WB83 ออนไลน์ พร้อมแผงส่ง secs_proxy_bot.exe และ config.ini ผ่าน TCP 5004")
    bullets(doc, [
        "Login ด้วย Role Developer และเลือก Machine ID ปลายทางให้ถูกต้อง",
        "ยืนยันว่า Deployment Receiver ทำงานบนเครื่องและ Port 5004 เปิดจาก RATS Server",
        "กด Choose Files แล้วเลือกเฉพาะ secs_proxy_bot.exe และ/หรือ config.ini",
        "เมื่อส่ง config.ini Backend จะเปลี่ยน machine_id ให้ตรงกับเครื่องที่เลือกโดยอัตโนมัติ",
        "ตรวจชื่อไฟล์ ห้ามเปลี่ยนชื่อและห้ามเลือกไฟล์อื่น",
        "กดส่งไฟล์ไปยังเครื่อง แล้วอ่าน Result: installed หรือ staged_pending",
        "หาก staged_pending ให้หยุด Recipe Bot บนเครื่อง เปลี่ยนไฟล์ .pending แทน EXE เดิม แล้ว Start Bot ใหม่",
        "ตรวจ Recipe Bot Status และ recipe_bot.log หลัง Deploy",
    ], numbered=True)

    heading(doc, "10. Authentication และ Role")
    table(doc, ["Role", "Level", "Backend action"], [
        ("Guest", "0", "ไม่มี Session"), ("Operator", "1", "Dashboard และ Pull/Push"),
        ("Technician", "2", "สิทธิ์ Operator รวม Delete และ Clear Log"),
        ("Administrator", "3", "ครอบคลุม Technician และดู Employee Audit"),
        ("Developer", "4", "ครอบคลุมทั้งหมดและ Deploy file"),
    ], [1.65, 0.75, 4.1])
    para(doc, "Login ต้องมี Employee Number ทุกครั้ง Backend ออก Token แบบ URL-safe 32 bytes เก็บใน Memory และต่ออายุเมื่อเรียก Protected API; ระบบบันทึก EN, Role, Login/Logout, เหตุผล timeout/restart และ Action ลง logs/employee_audit.json Frontend มี Inactivity warning ก่อนกลับ Guest")
    figure(doc, "15_developer_employee_audit.png", "รูปที่ D4 Employee Activity Audit สำหรับ Administrator/Developer", "ตาราง Audit แสดง Employee Number, Role, เวลา Login/Logout และ Action ที่ทำใน Session")

    heading(doc, "11. Recipe Workflow และการเก็บไฟล์")
    table(doc, ["กรณี", "Backend action"], [
        ("PPID ใหม่", "บันทึก .PWB ลง BondingProg ทันทีแบบ Atomic และตอบ 201 saved_new"),
        ("PPID เดิม + Hash เหมือน", "ตอบ identical และไม่สร้าง Popup/ไฟล์ซ้ำ"),
        ("PPID เดิม + Hash ต่าง", "Copy ไฟล์เดิมไป .archive พร้อม Timestamp แล้ว Replace ด้วยไฟล์ใหม่อัตโนมัติ"),
    ], [2.15, 4.35])
    para(doc, "Backend รองรับ PWB ที่ Gzip หรือไม่บีบอัด ใช้ Regex อ่าน Program Name และจำกัด PPID ไม่เกิน 120 ตัวอักษร พร้อมปฏิเสธอักขระที่ใช้เป็นชื่อไฟล์ไม่ได้")

    heading(doc, "12. Realtime State และ WebSocket")
    para(doc, "Frontend เปิด WebSocket ที่ /ws และรับ Snapshot ของ machines, events และ pending_recipe_updates Backend Broadcast เมื่อเกิดการเปลี่ยนจริง จึงไม่ต้อง Poll เครื่องตลอดเวลา")
    bullets(doc, [
        "machine_link_status มาจาก Worker/SECS-GEM",
        "bot_status มาจาก Host-side TCP 5003 Channel",
        "status เช่น IDLE/SYNCING/PUSHING/DELETING เป็นสถานะงานชั่วคราว",
        "WebSocket disconnect ทำให้ Frontend แสดง Backend Offline และพยายามเชื่อมใหม่ตามกลไก UI",
    ])

    heading(doc, "13. Start, Stop และ Log")
    bullets(doc, [
        "Start แบบ Production: ใช้ start_production.bat ใส่ Server IP, Build UI, Start Backend/Section Manager และ Serve Port 3000",
        "Start รายส่วน: python arc-system\\client-rats\\main.py; python arc-system\\section-manager\\manager.py; npm run preview -- --host 0.0.0.0 --port 3000 --strictPort",
        "Stop: stop_command_center.bat (ตรวจผล Process หลังใช้ เพราะ Script หยุด Python/Node หลาย Process)",
        "Machine log: watch_log.bat WB82 หรืออ่าน arc-system\\section-manager\\logs\\WB82.log",
        "Recipe Bot log: recipe_bot.log ข้าง EXE; Deployment Receiver ตรวจ Log/Console ตาม Build ที่ใช้งาน",
    ])
    heading(doc, "Procedure: เปิดระบบ Production", 2)
    bullets(doc, [
        "เปิด Command Prompt/PowerShell ด้วยสิทธิ์ตามนโยบายองค์กร",
        "รัน start_production.bat จาก Project root",
        "กรอก IPv4 ของ Server ที่ Operator PC ใช้เข้าถึง",
        "รอ Frontend build, Backend :8080 และ Section Manager เริ่มครบ",
        "เปิด http://<SERVER_IP>:3000 แล้วตรวจ Python Backend ONLINE",
        "Login เป็น Administrator และตรวจ Machine/Bot state; Login Developer เฉพาะเมื่อต้อง Deploy",
    ], numbered=True)
    figure(doc, "03_production_readiness_annotated.png", "รูปที่ D5 Production Readiness ที่อ่านสถานะจริง", "ภาพครอปหน้า Production Readiness พร้อมหมายเลขชี้ Overall, Section Manager, Recipe Storage และปุ่ม Refresh")
    table(doc, ["หมายเลข", "ตรวจอะไร", "เกณฑ์การตัดสินใจ"], [
        ("1", "Overall", "READY ทำงานได้; ATTENTION ต้องอ่าน Warning; NOT_READY ห้ามเริ่มงานสำคัญ"),
        ("2", "Section Manager", "ต้อง ONLINE และ Heartbeat ล่าสุดก่อนใช้ SECS/GEM"),
        ("3", "Storage/Worker", "ตรวจจำนวน Worker, Recipe, Archive, Audit และพื้นที่ Disk"),
        ("4", "Refresh", "กดหลัง Start/Restart หรือก่อน Acceptance test; หน้านี้ไม่ Poll เป็น loop"),
    ], [0.75, 1.65, 4.1])

    heading(doc, "14. Test และ Acceptance")
    callout(doc, "หลักฐานจากเครื่องจริง", "ภาพชุดนี้จัดทำขณะเชื่อมต่อเครื่องจริงและยืนยันเส้นทาง Recipe Bot จาก Popup Recipe Ready จนถึง Transfer Complete รวมทั้ง Dashboard Operator/Developer ที่เลือก WB#83 ออนไลน์", "ok")
    table(doc, ["Test", "เกณฑ์ผ่าน"], [
        ("Frontend/API", "เปิด :3000, Login ได้, Backend แสดง ONLINE และ WebSocket ไม่มี Error"),
        ("Machine link", "S1F2/Link test ผ่านและ Machine status เปลี่ยนแบบเรียลไทม์"),
        ("Bot link", "Backend Connect :5003, Token/Machine ID ผ่าน และ Bot status ONLINE"),
        ("New recipe", "Bot Popup → Accept → Host บันทึกด้วย PPID ภายใน PWB → Completion Popup"),
        ("Duplicate changed", "Backend Archive ไฟล์เดิมและ Overwrite ไฟล์ใหม่อัตโนมัติ โดยไม่รอ Approval"),
        ("Push", "PWB จาก Host ถูกส่ง SECS/GEM และ Recipe ปรากฏบนเครื่อง"),
        ("Delete", "S7F17/S7F18 สำเร็จและ Recipe หายจากเครื่อง"),
        ("Deploy", "Developer ส่ง allowlisted files; CRC ผ่าน; Locked EXE เป็น .pending"),
        ("Offline", "หยุด Bot/เครื่องแล้ว Dashboard เปลี่ยนเฉพาะสถานะที่เกี่ยวข้อง"),
    ], [1.75, 4.75])

    heading(doc, "15. Troubleshooting")
    table(doc, ["อาการ", "Root cause ที่พบบ่อย", "คำสั่ง/จุดตรวจ"], [
        ("Failed to fetch", "Browser เข้า Backend ไม่ได้/CORS/Port 8080", "Test-NetConnection <server> -Port 8080; ตรวจ RATS_CORS_ORIGINS"),
        ("Bot upload offline", "Port 5003/Token/Machine ID/Process", "Test-NetConnection <machine> -Port 5003; recipe_bot.log"),
        ("Machine online ไม่ขึ้น", "Port 5001/HSMS/Session/Worker", "Worker log และ SECS/GEM config"),
        ("Deploy rejected", "Token, AUTO machine validation, filename, CRC", "Receiver config + Backend Event Log"),
        ("EXE ไม่ถูกแทน", "Recipe Bot ล็อกไฟล์", "หยุด Bot, เปลี่ยน .pending เป็น secs_proxy_bot.exe, Start ใหม่"),
        ("Recipe ไม่ถูกอัปเดต", "PPID/Hash/Metadata ไม่ถูกต้อง หรือเขียน Archive ไม่สำเร็จ", "Event Log, BondingProg/.archive และ recipe_bot.log"),
        ("PPID not found", "ไม่มี Program Name ใน PWB หรือรูปแบบไม่รองรับ", "ตรวจ Raw/Decompressed PWB และ Regex"),
        ("Session 401", "หมดอายุ/Backend restart", "Login ใหม่และตรวจเวลาระบบ"),
    ], [1.55, 2.45, 2.5])

    heading(doc, "16. Security Hardening และ Change Control")
    bullets(doc, [
        "เปลี่ยน Default Password/Token และใช้ Named account/Secret store",
        "วาง Frontend/API หลัง HTTPS reverse proxy เมื่อออกนอก Trusted LAN",
        "จำกัด Firewall ตาม Source/Destination IP และเปิดเฉพาะ Port ที่ใช้",
        "สำรองและจำกัดสิทธิ์ logs/employee_audit.json ซึ่งบันทึก EN, Role, Machine, PPID, Result และ Timestamp",
        "สำรอง BondingProg/.archive และทดสอบ Restore",
        "Sign/Hash EXE ก่อน Deploy และเก็บ Release artifact แบบ Versioned",
        "ทดสอบกับเครื่อง Pilot เช่น WB#82 ก่อนกระจายทุกเครื่อง",
        "ทุกการเปลี่ยน database.py/config.ini ต้องมี Peer review, Rollback และ Acceptance test",
    ])
    callout(doc, "Release Gate", "ห้ามกระจาย Bot/Config ทั้ง Fleet จนกว่าจะผ่าน New/Changed/Identical recipe, Offline/Retry, Push/Delete และ Locked-EXE scenario บนเครื่อง Pilot", "danger")

    heading(doc, "17. Production Readiness และ Release Gate")
    para(doc, "ระบบที่ทำงานได้ยังไม่เท่ากับระบบพร้อมผลิตจริง ก่อนส่งมอบต้องมีหลักฐานว่าบริการหลัก สถานะเครื่อง ความปลอดภัย การสำรองข้อมูล และวิธีกู้คืนผ่านเกณฑ์เดียวกันทุกครั้ง หน้า System Status เรียก /api/health เมื่อเปิดหน้าและเมื่อกด Refresh โดยไม่สร้าง Polling loop ใหม่")
    figure(doc, "03_production_readiness_annotated.png", "รูปที่ D6 จุดตรวจ Production Readiness", "ภาพ Production Readiness พร้อมหมายเลขกำกับ Overall, Section Manager, Worker/Storage และ Refresh")
    table(doc, ["ระดับ", "ต้องทำก่อน Production"], [
        ("P0 - ห้ามข้าม", "เปลี่ยน Default password/token, ใช้ Production service แทน dev server, ทำ Backup/Restore test, ตั้ง Auto-start/Recovery และยืนยัน Section Manager heartbeat"),
        ("P1 - ควรครบ", "กำหนด Recipe retention/version, ตั้ง Log rotation, Export audit, Health alert, UPS/เวลาเครื่องตรงกัน และเอกสาร Rollback"),
        ("P2 - เพิ่มความเป็นระบบ", "Release version/build ID, Maintenance mode, Dashboard alarm acknowledgement, Fleet deployment history และรายงาน Acceptance ต่อรุ่น"),
    ], [1.45, 5.05])
    bullets(doc, [
        "ก่อนเริ่มกะ: Overall ต้องไม่เป็น NOT_READY และ Worker/Storage ที่เกี่ยวข้องต้อง READY",
        "ก่อน Push/Delete: ตรวจ Machine และ Bot แยกกัน พร้อมยืนยัน WB#/PPID และ Event Log ล่าสุด",
        "หลัง Deploy: บันทึก Build/Hash/ผู้ดำเนินการ/เครื่องเป้าหมาย และเก็บไฟล์ Rollback",
        "ทดสอบ Restore สูตรจาก .archive และ Audit backup ตามรอบที่หน่วยงานกำหนด",
        "หาก Credential Policy เป็น WARNING ต้องแก้ก่อนเปิดระบบให้ผู้ใช้หลายคนหรือหลายเครือข่าย",
    ])
    callout(doc, "สถานะปัจจุบัน", "ฟังก์ชันหลักพร้อมใช้งาน แต่ยังต้องปิด P0 โดยเฉพาะ Default credential, การรันเป็น Windows Service/Auto-recovery, Backup/Restore และ Production web serving ก่อนประกาศว่า Production Ready", "warn")
    return doc


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    user_path = OUT / "RATS_Command_Center_คู่มือผู้ใช้งาน_ภาษาไทย_ฉบับปรับปรุง.docx"
    dev_path = OUT / "RATS_Command_Center_คู่มือ_Developer_ภาษาไทย_ฉบับปรับปรุง.docx"
    user_manual().save(user_path)
    dev_manual().save(dev_path)
    print("Created 2 Thai DOCX manuals in docs/manuals")
